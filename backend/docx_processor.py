"""
DOCX processing module for extracting text from DOCX documents

Uses python-docx to read DOCX files and extract text content,
saving the results to output files for further processing.
"""

from docx import Document
import os
from datetime import datetime
from typing import Dict, Any, Optional
import logging
import io

logger = logging.getLogger(__name__)

class DOCXProcessor:
    def __init__(self, output_dir: str = "."):
        """
        Initialize DOCX processor with output directory
        
        Args:
            output_dir: Directory to save extracted text file (defaults to current directory)
        """
        self.output_dir = output_dir
        self.output_file = "design_doc.txt"
        # No need to create directory since we're using current directory
    
    def extract_text_from_docx(self, docx_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from DOCX content and save to single output file
        
        Args:
            docx_content: DOCX file content as bytes
            filename: Original filename of the DOCX
            
        Returns:
            Dictionary containing extraction results and metadata
        """
        try:
            # Use single output file path
            output_path = os.path.join(self.output_dir, self.output_file)
            
            # Extract text from DOCX
            extracted_text = ""
            paragraph_count = 0
            
            # Create a BytesIO object to work with python-docx
            docx_stream = io.BytesIO(docx_content)
            
            try:
                # Extract text using python-docx
                doc = Document(docx_stream)
                
                for paragraph in doc.paragraphs:
                    paragraph_text = paragraph.text.strip()
                    if paragraph_text:  # Only add non-empty paragraphs
                        extracted_text += paragraph_text + "\n\n"
                        paragraph_count += 1
                
                # Also extract text from tables
                table_count = 0
                for table in doc.tables:
                    table_count += 1
                    extracted_text += f"--- Table {table_count} ---\n"
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            extracted_text += " | ".join(row_text) + "\n"
                    extracted_text += "\n"
                
                # Write extracted text to output file
                with open(output_path, 'w', encoding='utf-8') as output_file:
                    output_file.write(f"DOCX Text Extraction Results\n")
                    output_file.write(f"Original File: {filename}\n")
                    output_file.write(f"Extracted: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    output_file.write(f"Total Paragraphs: {paragraph_count}\n")
                    output_file.write(f"Total Tables: {table_count}\n")
                    output_file.write("=" * 50 + "\n\n")
                    output_file.write(extracted_text)
                
                # Calculate statistics
                word_count = len(extracted_text.split()) if extracted_text.strip() else 0
                char_count = len(extracted_text)
                
                logger.info(f"Successfully extracted text from {filename}: {paragraph_count} paragraphs, {table_count} tables, {word_count} words")
                
                return {
                    "success": True,
                    "output_file": self.output_file,
                    "output_path": output_path,
                    "paragraph_count": paragraph_count,
                    "table_count": table_count,
                    "word_count": word_count,
                    "character_count": char_count,
                    "has_text": word_count > 0,
                    "extraction_time": datetime.now().isoformat()
                }
                
            except Exception as e:
                raise e
                
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {filename}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "output_file": None,
                "paragraph_count": 0,
                "table_count": 0,
                "word_count": 0,
                "character_count": 0,
                "has_text": False,
                "extraction_time": datetime.now().isoformat()
            }
    
    def get_extracted_text(self) -> Optional[str]:
        """
        Read content from the current extracted text file
        
        Returns:
            File content as string, or None if file doesn't exist
        """
        try:
            output_path = os.path.join(self.output_dir, self.output_file)
            if os.path.exists(output_path):
                with open(output_path, 'r', encoding='utf-8') as file:
                    return file.read()
            return None
        except Exception as e:
            logger.error(f"Error reading extracted text file: {str(e)}")
            return None
    
    def has_extracted_text(self) -> bool:
        """
        Check if extracted text file exists
        
        Returns:
            True if extracted text file exists, False otherwise
        """
        output_path = os.path.join(self.output_dir, self.output_file)
        return os.path.exists(output_path)


# Global DOCX processor instance
docx_processor = DOCXProcessor()